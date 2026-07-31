from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TRAVEL_TABLE.md"
OUTPUT = ROOT / "TRAVEL_TABLE.docx"


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shd = properties.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        properties.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    properties = tbl.tblPr

    tbl_width = properties.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        properties.append(tbl_width)
    tbl_width.set(qn("w:w"), str(sum(widths)))
    tbl_width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_keep_row_together(row):
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def parse_rows():
    rows = []
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        if not line.startswith("|") or line.startswith("|---"):
            continue
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        if values and values[0] != "Dates":
            rows.append(values)
    return rows


def set_run_font(run, size=8.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def main():
    headers = ["Dates", "Travel / route", "Type", "Status / booking"]
    source_rows = parse_rows()
    compact_rows = []
    for dates, travel, location, trip_type, status, booking in source_rows:
        compact_rows.append([dates, f"{travel}\n{location}", trip_type, f"{status}\n{booking}"])

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    widths = [1500, 4200, 1050, 2610]

    def add_table(rows, repeat_header=False):
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        set_table_geometry(table, widths)

        header = table.rows[0]
        if repeat_header:
            set_repeat_table_header(header)
        set_keep_row_together(header)
        for index, text in enumerate(headers):
            cell = header.cells[index]
            shade(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=8.5, bold=True, color="1F4D78")

        for row_values in rows:
            row = table.add_row()
            set_keep_row_together(row)
            for index, text in enumerate(row_values):
                cell = row.cells[index]
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for part_index, part in enumerate(text.split("\n")):
                    if part_index:
                        paragraph.add_run().add_break()
                    run = paragraph.add_run(part)
                    set_run_font(run, size=8.2, bold=(part_index == 0 and index in (0, 1, 2)))

        set_table_geometry(table, widths)

    # Split into two tables so the continuation carries an explicit header row.
    add_table(compact_rows[:23])
    add_table(compact_rows[23:], repeat_header=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
